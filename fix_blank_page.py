import docx

doc = docx.Document('templates/plantilla.docx')

for p in doc.paragraphs:
    # Set page break before ANEXO
    if p.text.strip() == "ANEXO":
        p.paragraph_format.page_break_before = True
    
    # Remove sectPr from paragraphs to avoid weird Word section blank pages
    if 'sectPr' in p._element.xml:
        for child in p._element:
            if child.tag.endswith('pPr'):
                for child2 in child:
                    if child2.tag.endswith('sectPr'):
                        child.remove(child2)
                        print("Removed a sectPr from a paragraph!")

# Now let's remove excessive empty paragraphs that might push it to a blank page
# Wait, actually we shouldn't delete elements while iterating
empty_streak = 0
paragraphs_to_delete = []

for p in doc.paragraphs:
    if not p.text.strip():
        empty_streak += 1
        if empty_streak > 1:
            # If there's more than 1 consecutive empty paragraph, delete it to avoid huge gaps
            paragraphs_to_delete.append(p)
    else:
        empty_streak = 0

for p in paragraphs_to_delete:
    p._element.getparent().remove(p._element)

doc.save('templates/plantilla.docx')
print("Fixed blank pages successfully.")
