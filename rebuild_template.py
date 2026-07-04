import docx
import re
import os

def process_document(input_path, output_path):
    doc = docx.Document(input_path)

    for p in doc.paragraphs:
        if not p.text.strip():
            continue
            
        text = p.text
        
        # 0. Nepotismo
        if "Yo," in text and "identificado (a) con DNI" in text:
            text = re.sub(r'Yo,\s*_+', 'Yo, {NOMBRE} ', text)
            text = re.sub(r'DNI N. \s*_+', 'DNI Nº {DNI} ', text)
            text = re.sub(r'domicilio en\s*_+', 'domicilio en {DIRECCION} ', text)

        # 1. First paragraph
        elif "El que suscribe" in text and "identificado con DNI" in text:
            text = re.sub(r'suscribe\s*_+\s*identificado', 'suscribe {NOMBRE} identificado', text)
            text = re.sub(r'DNI N.*?_+\s*con RUC', 'DNI N° {DNI} con RUC', text)
            text = re.sub(r'RUC N.*?_+\s*declaro', 'RUC N° {RUC} declaro', text)
            
        # 2. Second paragraph (Estudios)
        elif "El(la) que suscribe," in text and "estudios secundarios" in text:
            text = re.sub(r'suscribe,\s*_+\s*identificado', 'suscribe, {NOMBRE} identificado', text)
            text = re.sub(r'DNI N.*?_+\s*con RUC', 'DNI N° {DNI} con RUC', text)
            text = re.sub(r'RUC N.*?_+\s*domiciliado', 'RUC N° {RUC} domiciliado', text)
            text = re.sub(r'domiciliado\(a\) en\s*_+\s*declaro', 'domiciliado(a) en {DIRECCION} declaro', text)
            text = re.sub(r'centro de estudios\s*_+\s*en el a.o\s*_+', 'centro de estudios {COLEGIO} en el año {ANIO}', text)

        # 3. Third paragraph (Antecedentes)
        elif "El(la) que suscribe," in text and "ANTECEDENTES POLICIALES" in text:
            text = re.sub(r'suscribe,\s*_+\s*identificado', 'suscribe, {NOMBRE} identificado', text)
            text = re.sub(r'DNI N.*?\s*_+\s*, con RUC', 'DNI N° {DNI}, con RUC', text)
            text = re.sub(r'RUC N.*?\s*_+\s*, domiciliado', 'RUC N° {RUC}, domiciliado', text)
            text = re.sub(r'domiciliado\(a\) en\s*_+\s*, declaro', 'domiciliado(a) en {DIRECCION}, declaro', text)

        # 4. CCI
        elif "pagos a nombre de" in text and "sean abonados" in text:
            text = re.sub(r'nombre de\s*_+\s*, sean abonados', 'nombre de {NOMBRE}, sean abonados', text)
            text = re.sub(r'BANCO\s*_+', 'BANCO {BANCO}', text)
            
        # 5. Anexo Incompatibilidades
        elif "Yo" in text and "identificado con DNI" in text and "declaro bajo juramento" in text:
            text = re.sub(r'Yo\s*_+\s*identificado', 'Yo {NOMBRE} identificado', text)
            text = re.sub(r'DNI N.*?\s*_+', 'DNI N° {DNI}', text)

        # Signatures
        else:
            if "NOMBRES:" in text and not "{NOMBRE}" in text:
                text = text.replace("NOMBRES:", "NOMBRES: {NOMBRE}")
            if "NOMBRE:" in text and not "{NOMBRE}" in text and not "NOMBRES" in text:
                text = text.replace("NOMBRE:", "NOMBRE: {NOMBRE}")
            
            # Using regex to replace DNI but preserve leading spaces
            if re.search(r'DNI N[°º]?\s*$', text) or re.search(r'DNI N[°º]?\s*:\s*$', text):
                text = re.sub(r'(DNI N[°º]?\s*:?)\s*$', r'\1 {DNI}', text)
            elif re.search(r'DNI\s*:\s*$', text):
                text = re.sub(r'(DNI\s*:)\s*$', r'\1 {DNI}', text)
            
            if re.search(r'RUC N[°º]?\s*$', text) or re.search(r'RUC N[°º]?\s*:\s*$', text):
                text = re.sub(r'(RUC N[°º]?\s*:?)\s*$', r'\1 {RUC}', text)
            elif re.search(r'RUC\s*:\s*$', text):
                text = re.sub(r'(RUC\s*:)\s*$', r'\1 {RUC}', text)

        if text != p.text:
            p.text = text

    # Add tags to CCI table (Table 1)
    if len(doc.tables) > 1:
        cci_table = doc.tables[1]
        if len(cci_table.rows) > 0 and len(cci_table.columns) == 20:
            for i, cell in enumerate(cci_table.rows[0].cells):
                cell.text = f"{{c{i}}}"

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Template successfully saved to {output_path}")

input_file = "C:/Users/Soporte/Desktop/Cecom-20251224T131857Z-3-002/Cecom/PAGOS/DJJunio-mano.docx"
output_file = "C:/Users/Soporte/Desktop/Cecom-20251224T131857Z-3-002/Cecom/BACKEND/sistema_pagos_dj/templates/plantilla.docx"
process_document(input_file, output_file)
